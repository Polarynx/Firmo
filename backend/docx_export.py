"""The paper, as a Word document.

Firmo can hold the whole research process, but almost nobody submits from it —
the file that gets graded is a .docx. Exporting plain text would throw away the
part that costs students marks: a works-cited page with real hanging indents,
correctly ordered, with journal titles actually italicised.

So the export is a real document. Body copy double-spaced in the discipline's
usual face, the reference list on its own page under the heading that style
calls for, and the italics carried through from the publisher's own record.
"""
import html
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

# What each style calls the list at the back. Getting this wrong is the kind of
# detail a marker notices immediately.
HEADING = {
    "apa": "References",
    "mla": "Works Cited",
    "chicago": "Bibliography",
    "harvard": "Reference list",
    "ieee": "References",
}

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)

# The subset of markup CrossRef puts in a formatted citation.
_TAG_RE = re.compile(r"<(/?)(i|em|b|strong|sup|sub)\b[^>]*>", re.I)


def _segments(markup: str) -> list[tuple[str, bool, bool]]:
    """Split citation markup into (text, italic, bold) runs.

    CrossRef returns `<i>Nature</i>` on the exact-citation path, so a naive
    strip would leave the journal title unformatted and a naive insert would
    print the tags. Anything outside the small tag set is treated as text.
    """
    out: list[tuple[str, bool, bool]] = []
    italic = 0
    bold = 0
    pos = 0
    for m in _TAG_RE.finditer(markup or ""):
        if m.start() > pos:
            out.append((html.unescape(markup[pos:m.start()]), italic > 0, bold > 0))
        closing, tag = m.group(1) == "/", m.group(2).lower()
        delta = -1 if closing else 1
        if tag in ("i", "em"):
            italic = max(0, italic + delta)
        elif tag in ("b", "strong"):
            bold = max(0, bold + delta)
        pos = m.end()
    if pos < len(markup or ""):
        out.append((html.unescape(markup[pos:]), italic > 0, bold > 0))
    return [s for s in out if s[0]]


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    # Word ignores double-spacing set only on the style's paragraph format for
    # some templates, so it is set per paragraph as well below.
    normal.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def _body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    return p


def build_docx(text: str, entries: list[dict], style: str = "apa",
               title: str = "", author: str = "") -> bytes:
    """Assemble the document. `entries` come from citations.format_bibliography."""
    style = (style or "apa").lower()
    doc = Document()
    _style_document(doc)

    # MLA puts the student's name at the top left; the other styles centre a
    # title. Both are omitted entirely when the student gave us neither, rather
    # than printing an empty heading.
    if author and style == "mla":
        for line in author.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.add_run(line.strip())

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = style in ("apa", "harvard", "ieee")
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE

    if author and style != "mla":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        p.add_run(author)

    # The draft. Blank lines are paragraph breaks, which is exactly how the
    # canvas treats them, so what the student sees is what they get.
    body = (text or "").replace("\r\n", "\n").strip()
    if body:
        for block in re.split(r"\n\s*\n", body):
            block = block.strip()
            if block:
                _body_paragraph(doc, block)

    if not entries:
        return _to_bytes(doc)

    # The reference list starts its own page.
    if body or title or author:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.line_spacing = 2.0
    h.paragraph_format.space_after = Pt(0)
    hr = h.add_run(HEADING.get(style, "References"))
    hr.bold = True
    hr.font.name = BODY_FONT
    hr.font.size = BODY_SIZE

    for entry in entries:
        citation = entry.get("citation") or ""
        if not citation.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        # A hanging indent is a left indent with a negative first line, which
        # is what every style guide means by "indent the second line onward".
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for chunk, italic, bold in _segments(citation):
            run = p.add_run(chunk)
            run.italic = italic
            run.bold = bold
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE

    return _to_bytes(doc)


def _to_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(name: str, fallback: str = "paper") -> str:
    """A filename Windows, macOS, and Linux will all accept."""
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", (name or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(". ")
    return (cleaned or fallback)[:80]
